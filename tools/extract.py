"""Extract text + image positions from docx, preserving document order."""
import os, zipfile, json
from docx import Document
from docx.oxml.ns import qn
import xml.etree.ElementTree as ET

DOCX = "tools/source.docx"
OUT_DIR = "tools/extracted"
os.makedirs(OUT_DIR, exist_ok=True)

doc = Document(DOCX)

with zipfile.ZipFile(DOCX) as z:
    rels_xml = z.read("word/_rels/document.xml.rels").decode("utf-8")
ns = {"r": "http://schemas.openxmlformats.org/package/2006/relationships"}
rels_root = ET.fromstring(rels_xml)
rid_to_target = {}
for rel in rels_root.findall("r:Relationship", ns):
    rid_to_target[rel.get("Id")] = rel.get("Target")

output = []
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    images = []
    for blip in para._p.iter(qn("a:blip")):
        rid = blip.get(qn("r:embed")) or blip.get(qn("r:link"))
        if rid and rid in rid_to_target:
            images.append(rid_to_target[rid])
    style = para.style.name if para.style else ""
    if text or images:
        output.append({"idx": i, "style": style, "text": text, "images": images})

# Also walk tables
table_data = []
for ti, table in enumerate(doc.tables):
    rows = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        rows.append(cells)
    table_data.append({"idx": ti, "rows": rows})

with open(os.path.join(OUT_DIR, "doc_order.json"), "w", encoding="utf-8") as f:
    json.dump({"paragraphs": output, "tables": table_data}, f, ensure_ascii=False, indent=2)

print(f"Total paragraphs: {len(output)}")
print(f"Total tables: {len(table_data)}")
print("="*80)
for entry in output:
    if entry["images"]:
        print(f"[{entry['idx']:03d}] [{entry['style']}] >>>IMG: {entry['images']}")
    if entry["text"]:
        print(f"[{entry['idx']:03d}] [{entry['style']}] {entry['text'][:300]}")
print("="*80)
print("TABLES:")
for t in table_data:
    print(f"--- Table {t['idx']} ---")
    for r in t["rows"]:
        print(" | ".join(r))
